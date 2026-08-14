import os
import re

from pypdf import PdfReader
from docx import Document


class ScriptParseError(Exception):
    pass


SCENE_HEADING_RE = re.compile(
    r"""
    ^
    \s*
    (?:\d+[A-Z]?\s+)?                         # Optional scene number: 12 or 12A
    (?:
        INT(?:ERIOR)?|EXT(?:ERIOR)?|           # INT, INTERIOR, EXT, EXTERIOR
        INT\s*/\s*EXT|EXT\s*/\s*INT|           # INT/EXT, EXT/INT
        I\s*/\s*E|E\s*/\s*I                   # I/E, E/I
    )
    \s*
    (?:[./:;-]|\s+-\s+|\s+)
    .+
    $
    """,
    re.IGNORECASE | re.VERBOSE,
)


def extract_text(stored_path: str) -> str:
    """
    Extract plain text from a screenplay file on disk.
    Dispatches based on file extension. Raises ScriptParseError on failure.
    """
    _, ext = os.path.splitext(stored_path)
    ext = ext.lower()

    if ext == ".pdf":
        text = _extract_from_pdf(stored_path)
    elif ext in (".doc", ".docx"):
        text = _extract_from_docx(stored_path)
    else:
        raise ScriptParseError(f"Unsupported file type for parsing: {ext}")

    text = text.strip()
    if not text:
        raise ScriptParseError(
            "No extractable text was found in this file. It may be a scanned "
            "image-only PDF, which isn't supported yet."
        )
    return text


def _extract_from_pdf(path: str) -> str:
    try:
        reader = PdfReader(path)
    except Exception as e:
        raise ScriptParseError(f"Could not open PDF: {e}") from e

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            # Skip pages that fail to extract rather than failing the whole script
            continue

    return "\n\n".join(pages_text)


def _extract_from_docx(path: str) -> str:
    try:
        document = Document(path)
    except Exception as e:
        raise ScriptParseError(f"Could not open Word document: {e}") from e

    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of any tables (some script templates use them for scene headers)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def basic_scene_count(raw_text: str) -> int:
    """
    Rough heuristic scene counter for common screenplay slugline variants,
    useful as a quick sanity signal before the real AI breakdown runs.
    """
    return sum(
        1
        for line in raw_text.splitlines()
        if SCENE_HEADING_RE.match(line.strip())
    )
