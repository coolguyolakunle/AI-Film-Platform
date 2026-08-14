import io
import tempfile
from unittest.mock import patch

import pytest

from app import create_app
from app.extensions import db
from app.services.breakdown_schema import BreakdownOutput

TEST_EMAIL = "export.tester" + "@" + "filmtest.local"

SAMPLE_BREAKDOWN_JSON = {
    "scenes": [
        {
            "scene_number": 1,
            "heading": "INT. WAREHOUSE - NIGHT",
            "int_ext": "INT",
            "time_of_day": "NIGHT",
            "location": "Warehouse",
            "synopsis": "John searches the warehouse.",
            "characters": ["JOHN"],
            "props": ["flashlight"],
            "costumes": [],
            "departments": ["Camera"],
        }
    ],
    "all_characters": ["JOHN"],
    "all_locations": ["Warehouse"],
    "all_props": ["flashlight"],
    "all_costumes": [],
}


@pytest.fixture
def app():
    upload_dir = tempfile.mkdtemp()
    app = create_app("testing")
    app.config["UPLOAD_FOLDER"] = upload_dir
    with app.app_context():
        db.create_all()
        yield app
        db.session.remove()
        db.drop_all()


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def auth_headers(client):
    resp = client.post("/api/auth/register", json={
        "name": "Export Tester", "email": TEST_EMAIL, "password": "password123",
    })
    token = resp.get_json()["token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def script_with_breakdown(client, auth_headers):
    proj_resp = client.post("/api/projects", json={"title": "Export Test Film"}, headers=auth_headers)
    project_id = proj_resp.get_json()["project"]["id"]

    from docx import Document
    doc = Document()
    doc.add_paragraph("INT. WAREHOUSE - NIGHT")
    doc.add_paragraph("JOHN searches with a flashlight.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    upload_resp = client.post(
        "/api/scripts/upload",
        data={"project_id": project_id, "file": (buf, "my script.docx")},
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    script_id = upload_resp.get_json()["script"]["id"]

    with patch("app.tasks.analyze_script") as mock_analyze:
        mock_analyze.return_value = BreakdownOutput.model_validate(SAMPLE_BREAKDOWN_JSON)
        client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)

    return script_id


def test_export_pdf(client, auth_headers, script_with_breakdown):
    resp = client.get(
        f"/api/breakdowns/{script_with_breakdown}/export?format=pdf", headers=auth_headers
    )
    assert resp.status_code == 200
    assert resp.mimetype == "application/pdf"
    assert resp.data[:4] == b"%PDF"
    assert "attachment" in resp.headers["Content-Disposition"]


def test_export_csv(client, auth_headers, script_with_breakdown):
    resp = client.get(
        f"/api/breakdowns/{script_with_breakdown}/export?format=csv", headers=auth_headers
    )
    assert resp.status_code == 200
    assert resp.mimetype == "text/csv"
    body = resp.data.decode("utf-8")
    assert "Scene #" in body
    assert "WAREHOUSE" in body


def test_export_defaults_to_pdf(client, auth_headers, script_with_breakdown):
    resp = client.get(f"/api/breakdowns/{script_with_breakdown}/export", headers=auth_headers)
    assert resp.status_code == 200
    assert resp.mimetype == "application/pdf"


def test_export_rejects_bad_format(client, auth_headers, script_with_breakdown):
    resp = client.get(
        f"/api/breakdowns/{script_with_breakdown}/export?format=docx", headers=auth_headers
    )
    assert resp.status_code == 400
    assert resp.get_json()["error"] == "invalid_format"


def test_export_before_breakdown_ready_422s(client, auth_headers):
    proj_resp = client.post("/api/projects", json={"title": "No Breakdown Yet"}, headers=auth_headers)
    project_id = proj_resp.get_json()["project"]["id"]

    from docx import Document
    doc = Document()
    doc.add_paragraph("INT. ROOM - DAY")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    upload_resp = client.post(
        "/api/scripts/upload",
        data={"project_id": project_id, "file": (buf, "no_breakdown.docx")},
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    script_id = upload_resp.get_json()["script"]["id"]

    resp = client.get(f"/api/breakdowns/{script_id}/export", headers=auth_headers)
    assert resp.status_code == 422


def test_export_requires_ownership(client, auth_headers, script_with_breakdown):
    other_email = "other.export" + "@" + "filmtest.local"
    reg = client.post("/api/auth/register", json={
        "name": "Other", "email": other_email, "password": "password123",
    })
    other_headers = {"Authorization": f"Bearer {reg.get_json()['token']}"}

    resp = client.get(
        f"/api/breakdowns/{script_with_breakdown}/export", headers=other_headers
    )
    assert resp.status_code == 404
