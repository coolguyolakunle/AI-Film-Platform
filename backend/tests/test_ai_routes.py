import io
import tempfile
from unittest.mock import patch

import pytest
from kombu.exceptions import OperationalError

from app import create_app
from app.extensions import db
from app.services.breakdown_schema import BreakdownOutput
from app.services.ai_service import AIServiceError

TEST_EMAIL = "ai.tester" + "@" + "filmtest.local"

SAMPLE_BREAKDOWN_JSON = {
    "scenes": [
        {
            "scene_number": 1,
            "heading": "INT. WAREHOUSE - NIGHT",
            "int_ext": "INT",
            "time_of_day": "NIGHT",
            "location": "Warehouse",
            "synopsis": "John searches the warehouse.",
            "characters": ["JOHN"],
            "props": ["flashlight"],
            "costumes": [],
            "departments": ["Camera"],
        }
    ],
    "all_characters": ["JOHN"],
    "all_locations": ["Warehouse"],
    "all_props": ["flashlight"],
    "all_costumes": [],
}


@pytest.fixture
def app():
    upload_dir = tempfile.mkdtemp()
    app = create_app("testing")
    app.config["UPLOAD_FOLDER"] = upload_dir
    with app.app_context():
        db.create_all()
        yield app
        db.session.remove()
        db.drop_all()


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def auth_headers(client):
    resp = client.post("/api/auth/register", json={
        "name": "AI Tester", "email": TEST_EMAIL, "password": "password123",
    })
    token = resp.get_json()["token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def script_id(client, auth_headers):
    proj_resp = client.post("/api/projects", json={"title": "AI Test Film"}, headers=auth_headers)
    project_id = proj_resp.get_json()["project"]["id"]

    from docx import Document
    doc = Document()
    doc.add_paragraph("INT. WAREHOUSE - NIGHT")
    doc.add_paragraph("JOHN searches with a flashlight.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    upload_resp = client.post(
        "/api/scripts/upload",
        data={"project_id": project_id, "file": (buf, "script.docx")},
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    return upload_resp.get_json()["script"]["id"]


# NOTE: CELERY_TASK_ALWAYS_EAGER=True in the testing config means .delay()
# runs the task inline, synchronously, in the same request. That's the
# standard way to test Celery-backed code without a real broker/worker —
# it means these tests see the *final* state (complete/failed) in the
# response, whereas in production the request returns 202 immediately and
# the frontend polls GET /breakdowns/<script_id> until the worker finishes.
# A real live-worker run (outside pytest) is used separately to prove the
# actual async behavior end-to-end.


def test_analyze_returns_202_and_enqueues(client, auth_headers, script_id):
    with patch("app.tasks.analyze_script") as mock_analyze:
        mock_analyze.return_value = BreakdownOutput.model_validate(SAMPLE_BREAKDOWN_JSON)

        resp = client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)

        assert resp.status_code == 202
        body = resp.get_json()
        # Eager mode means the task already completed by the time we get a response.
        assert body["breakdown"]["status"] == "complete"
        assert body["breakdown"]["ai_output_json"]["summary"]["total_scenes"] == 1


def test_analyze_returns_503_when_queue_unavailable(client, auth_headers, script_id):
    with patch("app.routes.ai_routes.run_script_analysis.delay") as mock_delay:
        mock_delay.side_effect = OperationalError("Redis unavailable")

        resp = client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)

    assert resp.status_code == 503
    body = resp.get_json()
    assert body["error"] == "analysis_queue_unavailable"

    script_resp = client.get(f"/api/scripts/{script_id}", headers=auth_headers)
    assert script_resp.get_json()["script"]["status"] == "parsed"

    breakdown_resp = client.get(f"/api/breakdowns/{script_id}", headers=auth_headers)
    assert breakdown_resp.status_code == 404


def test_analyze_transient_ai_error_does_not_crash_request(client, auth_headers, script_id):
    # With real retries available (max_retries=2, the production default),
    # a single AIServiceError should be absorbed into a scheduled retry
    # rather than crashing the request or the eager-mode task run. Eager
    # mode has no real scheduler, so it can't simulate the full multi-attempt
    # retry-then-fail cycle here — that's verified separately with a live
    # Redis + worker run. This test proves the request stays healthy and the
    # script is left in "processing" (not silently stuck in some broken
    # state) when a transient failure occurs.
    with patch("app.tasks.analyze_script") as mock_analyze:
        mock_analyze.side_effect = AIServiceError("AI provider request failed: timeout")
        resp = client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)
        assert resp.status_code == 202

    get_resp = client.get(f"/api/scripts/{script_id}", headers=auth_headers)
    assert get_resp.get_json()["script"]["status"] == "processing"


def test_analyze_requires_ownership(client, auth_headers, script_id):
    other_email = "other.tester" + "@" + "filmtest.local"
    reg = client.post("/api/auth/register", json={
        "name": "Other", "email": other_email, "password": "password123",
    })
    other_headers = {"Authorization": f"Bearer {reg.get_json()['token']}"}

    resp = client.post(f"/api/ai/analyze/{script_id}", headers=other_headers)
    assert resp.status_code == 404


def test_analyze_rejects_double_submission(client, auth_headers, script_id):
    # Manually put the script in "processing" to simulate a job already in flight
    from app.extensions import db
    from app.models.script import Script
    script = db.session.get(Script, script_id)
    script.status = "processing"
    db.session.commit()

    resp = client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)
    assert resp.status_code == 409
    assert resp.get_json()["error"] == "already_processing"


def test_get_breakdown_before_analysis_404s(client, auth_headers, script_id):
    resp = client.get(f"/api/breakdowns/{script_id}", headers=auth_headers)
    assert resp.status_code == 404


def test_get_breakdown_after_analysis(client, auth_headers, script_id):
    with patch("app.tasks.analyze_script") as mock_analyze:
        mock_analyze.return_value = BreakdownOutput.model_validate(SAMPLE_BREAKDOWN_JSON)
        client.post(f"/api/ai/analyze/{script_id}", headers=auth_headers)

    resp = client.get(f"/api/breakdowns/{script_id}", headers=auth_headers)
    assert resp.status_code == 200
    assert resp.get_json()["breakdown"]["status"] == "complete"
