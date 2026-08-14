import io
import os
import tempfile

import pytest

from app import create_app
from app.extensions import db
from app.services.script_parser import basic_scene_count

TEST_EMAIL = "script.tester" + "@" + "filmtest.local"


@pytest.fixture
def app():
    upload_dir = tempfile.mkdtemp()
    app = create_app("testing")
    app.config["UPLOAD_FOLDER"] = upload_dir
    with app.app_context():
        db.create_all()
        yield app
        db.session.remove()
        db.drop_all()


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def auth_headers(client):
    resp = client.post("/api/auth/register", json={
        "name": "Script Tester", "email": TEST_EMAIL, "password": "password123",
    })
    token = resp.get_json()["token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
def project_id(client, auth_headers):
    resp = client.post("/api/projects", json={"title": "Test Film"}, headers=auth_headers)
    return resp.get_json()["project"]["id"]


def _sample_docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_paragraph("INT. WAREHOUSE - NIGHT")
    doc.add_paragraph("JOHN enters holding a FLASHLIGHT.")
    doc.add_paragraph("EXT. STREET - DAY")
    doc.add_paragraph("MARY walks by.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def test_upload_docx_success(client, auth_headers, project_id):
    file_bytes = _sample_docx_bytes()
    data = {
        "project_id": project_id,
        "file": (io.BytesIO(file_bytes), "script.docx"),
    }
    resp = client.post(
        "/api/scripts/upload",
        data=data,
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    assert resp.status_code == 201
    body = resp.get_json()
    assert body["script"]["status"] == "parsed"
    assert body["scene_count_estimate"] == 2


def test_basic_scene_count_detects_real_world_slugline_variants():
    raw_text = """
    INT. WAREHOUSE - NIGHT
    EXT: STREET - DAY
    INT - KITCHEN - MORNING
    INT/EXT. MOVING CAR - CONTINUOUS
    I/E. MOTEL ROOM / PARKING LOT - NIGHT
    INTERIOR OFFICE DAY
    EXTERIOR ROOFTOP - SUNSET
    12A EXT. ALLEY - NIGHT
    """

    assert basic_scene_count(raw_text) == 8


def test_basic_scene_count_avoids_common_false_positives():
    raw_text = """
    JOHN
    I expect this to work.

    The exterior wall is cracked.
    The interior light flickers.
    EXTREME CLOSE UP ON THE KNIFE
    INTENTIONALLY, she says nothing.
    """

    assert basic_scene_count(raw_text) == 0


def test_upload_rejects_bad_extension(client, auth_headers, project_id):
    data = {
        "project_id": project_id,
        "file": (io.BytesIO(b"not a real script"), "script.txt"),
    }
    resp = client.post(
        "/api/scripts/upload",
        data=data,
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    assert resp.status_code == 400
    assert resp.get_json()["error"] == "invalid_file"


def test_upload_requires_project_ownership(client, auth_headers):
    data = {
        "project_id": "nonexistent-project-id",
        "file": (io.BytesIO(_sample_docx_bytes()), "script.docx"),
    }
    resp = client.post(
        "/api/scripts/upload",
        data=data,
        headers=auth_headers,
        content_type="multipart/form-data",
    )
    assert resp.status_code == 404


def test_upload_requires_auth(client, project_id):
    data = {
        "project_id": project_id,
        "file": (io.BytesIO(_sample_docx_bytes()), "script.docx"),
    }
    resp = client.post(
        "/api/scripts/upload",
        data=data,
        content_type="multipart/form-data",
    )
    assert resp.status_code == 401


def test_list_scripts_for_project(client, auth_headers, project_id):
    data = {
        "project_id": project_id,
        "file": (io.BytesIO(_sample_docx_bytes()), "script.docx"),
    }
    client.post("/api/scripts/upload", data=data, headers=auth_headers, content_type="multipart/form-data")

    resp = client.get(f"/api/scripts/project/{project_id}", headers=auth_headers)
    assert resp.status_code == 200
    scripts = resp.get_json()["scripts"]
    assert len(scripts) == 1
    assert scripts[0]["original_filename"] == "script.docx"


def test_delete_script(client, auth_headers, project_id):
    data = {
        "project_id": project_id,
        "file": (io.BytesIO(_sample_docx_bytes()), "script.docx"),
    }
    upload_resp = client.post(
        "/api/scripts/upload", data=data, headers=auth_headers, content_type="multipart/form-data"
    )
    script_id = upload_resp.get_json()["script"]["id"]

    del_resp = client.delete(f"/api/scripts/{script_id}", headers=auth_headers)
    assert del_resp.status_code == 200

    get_resp = client.get(f"/api/scripts/{script_id}", headers=auth_headers)
    assert get_resp.status_code == 404
